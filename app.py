import streamlit as st
import random
import requests
from xml.etree import ElementTree
from io import BytesIO
from docx import Document

PUBMED_API_KEY = "47c558837273db73214cd0778f0ff099c908"
HUGGINGFACE_SUMMARY_API = "https://api-inference.huggingface.co/models/facebook/bart-large-cnn"

st.title("Medical Research Prompt Bot - End-to-End Research Assistant")

# Summarization using HuggingFace Inference API
def summarize_text(text):
    headers = {"Authorization": "Bearer YOUR_HUGGINGFACE_API_KEY"}
    payload = {"inputs": text, "options": {"min_length": 50, "max_length": 100}}
    response = requests.post(HUGGINGFACE_SUMMARY_API, headers=headers, json=payload)
    if response.status_code == 200:
        summary = response.json()[0]['summary_text']
        return summary
    else:
        return "Summarization failed."

# Research Question Generator
def generate_research_question():
    conditions = ["chronic kidney disease", "heart failure", "diabetes mellitus", "hypertension", "sepsis"]
    interventions = ["SGLT2 inhibitors", "ACE inhibitors", "beta-blockers", "early goal-directed therapy", "continuous renal replacement therapy"]
    outcomes = ["mortality reduction", "hospitalization rates", "renal function preservation", "cardiac remodeling", "infection control"]

    population = f"Patients with {random.choice(conditions)}"
    intervention = random.choice(interventions)
    outcome = random.choice(outcomes)

    question = f"What is the efficacy of {intervention} in {population} regarding {outcome}?"
    return population, intervention, "Placebo or standard care", outcome, question

if st.button("Generate Random Research Question"):
    population, intervention, comparison, outcome, research_topic = generate_research_question()
    st.session_state["population"] = population
    st.session_state["intervention"] = intervention
    st.session_state["comparison"] = comparison
    st.session_state["outcome"] = outcome
    st.session_state["research_topic"] = research_topic

research_topic = st.text_input("Research Topic", st.session_state.get("research_topic", ""))
population = st.text_input("Population (P)", st.session_state.get("population", ""))
intervention = st.text_input("Intervention (I)", st.session_state.get("intervention", ""))
comparison = st.text_input("Comparison (C)", st.session_state.get("comparison", ""))
outcome = st.text_input("Outcome (O)", st.session_state.get("outcome", ""))
study_types = st.multiselect("Study Types", ["systematic reviews", "meta-analyses", "RCTs"], default=["systematic reviews", "meta-analyses", "RCTs"])
publication_year_limit = st.slider("Publication Year Limit (years)", 1, 10, 5)
preferred_sources = st.multiselect("Preferred Sources", ["PubMed", "Cochrane Library", "Guidelines"], default=["PubMed", "Cochrane Library", "Guidelines"])
output_format = st.selectbox("Output Format", ["bullet points", "detailed report", "table"])

# Generate PubMed Search Query
def generate_pubmed_query(population, intervention, comparison, outcome):
    query_parts = []
    if population:
        query_parts.append(f"{population}")
    if intervention:
        query_parts.append(f"{intervention}")
    if comparison:
        query_parts.append(f"{comparison}")
    if outcome:
        query_parts.append(f"{outcome}")
    return " AND ".join(query_parts)

# Fetch PubMed Articles
def fetch_pubmed_articles(query, retmax=5):
    base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi"
    params = {
        'db': 'pubmed',
        'term': query,
        'retmode': 'xml',
        'retmax': retmax,
        'api_key': PUBMED_API_KEY
    }
    response = requests.get(base_url, params=params)
    root = ElementTree.fromstring(response.content)
    ids = [id_elem.text for id_elem in root.findall(".//Id")]
    return ids

# Fetch Article Details (Title, Abstract, DOI)
def fetch_pubmed_details(id_list):
    details = []
    if not id_list:
        return details
    base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi"
    params = {
        'db': 'pubmed',
        'id': ','.join(id_list),
        'retmode': 'xml',
        'api_key': PUBMED_API_KEY
    }
    response = requests.get(base_url, params=params)
    root = ElementTree.fromstring(response.content)

    for article in root.findall(".//PubmedArticle"):
        title_elem = article.find(".//ArticleTitle")
        abstract_elem = article.find(".//AbstractText")
        article_id_elem = article.find(".//ArticleId[@IdType='doi']")
        
        title = title_elem.text if title_elem is not None else "No Title"
        abstract = abstract_elem.text if abstract_elem is not None else "No Abstract"
        doi = article_id_elem.text if article_id_elem is not None else ""

        details.append({
            "title": title,
            "abstract": abstract,
            "doi": doi
        })
    return details

# Generate Full Report Document
def generate_docx_report(prompt, articles):
    doc = Document()
    doc.add_heading("Medical Research Prompt Report", level=1)
    doc.add_paragraph(prompt)
    doc.add_heading("Articles and Summaries", level=2)
    for idx, article in enumerate(articles, 1):
        doc.add_heading(f"{idx}. {article['title']}", level=3)
        doc.add_paragraph(f"Abstract: {article['abstract']}")
        if article['doi']:
            doc.add_paragraph(f"DOI: https://doi.org/{article['doi']}")
        if article['abstract'] != "No Abstract":
            summary = summarize_text(article['abstract'])
            doc.add_paragraph(f"Summary: {summary}")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

if st.button("Generate Prompt & Search PubMed"):
    prompt = f"""
Act as a clinical researcher specialized in evidence-based medicine. I need you to conduct a structured literature review on the following topic:

**Research Topic:** {research_topic}

Use the PICO framework:
- **Population (P):** {population}
- **Intervention (I):** {intervention}
- **Comparison (C):** {comparison}
- **Outcome (O):** {outcome}

**Instructions:**
1. Retrieve evidence from the following study types: {', '.join(study_types)}.
2. Only include studies published within the last {publication_year_limit} years.
3. Prefer references from: {', '.join(preferred_sources)}.
4. Provide references with PubMed links or DOI.
5. Present the findings in {output_format}.
6. Summarize key conclusions and highlight gaps in current research.

Optional: If applicable, suggest potential future research directions.
"""

    pubmed_query = generate_pubmed_query(population, intervention, comparison, outcome)

    st.subheader("Generated Research Prompt")
    st.text_area("Prompt", prompt, height=400)

    st.subheader("PubMed Search Query")
    st.code(pubmed_query, language="text")

    st.subheader("Fetching Articles from PubMed...")
    ids = fetch_pubmed_articles(pubmed_query)
    details = fetch_pubmed_details(ids)

    if details:
        for idx, article in enumerate(details, 1):
            st.markdown(f"**{idx}. {article['title']}**")
            st.markdown(f"Abstract: {article['abstract']}")
            if article['doi']:
                st.markdown(f"[Link to Article](https://doi.org/{article['doi']})")

            if article['abstract'] != "No Abstract":
                summary = summarize_text(article['abstract'])
                st.markdown(f"**Summary:** {summary}")

            st.markdown("---")

        report_file = generate_docx_report(prompt, details)
        st.download_button(
            label="Download Full Report (.docx)",
            data=report_file,
            file_name="Research_Prompt_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.write("No articles found.")

st.markdown("---")
st.markdown("Developed by Khalid Al — Medical Research Prompt Bot")
